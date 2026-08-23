"""DOCX exporter for rendering WorkingDocuments as styled Word files.

Renders onto a fresh in-memory copy of the real Colvir corporate template
(``formatting/templates/colvir.docx``, built by
``tools/build_default_docx_template.py`` from the user-supplied
``colvir_v3.dotx``). Every render loads the base asset from disk and never
writes back to it -- the shared template file is never mutated in place.
"""

from __future__ import annotations

import re
from collections.abc import Iterator
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

import docx
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt

from docgen.documents.schemas import DocumentNode, NodeKind, WorkingDocument
from docgen.export._naming import make_safe_filename
from docgen.export.html import ImageAsset, ImageLoader, local_storage_image_loader
from docgen.export.protocol import RenderedFile
from docgen.formatting.schemas import FormattingTemplate
from docgen.templates_catalog.use_case import (
    USE_CASE_DESCRIPTION_ROWS,
    USE_CASE_HISTORY_HEADERS,
    USE_CASE_LINK_HEADERS,
    USE_CASE_METADATA_FIELDS,
    USE_CASE_TECHNICAL_HEADERS,
)

__all__ = [
    "DocxExporter",
    "ImageAsset",
    "ImageLoader",
    "document_category_label",
    "filename_title",
    "local_storage_image_loader",
]

_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "formatting" / "templates"

_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Human labels for the semantic build templates (see
# templates_catalog/semantic/*.yaml's own `name:`), used to swap the
# corporate footer's generic "Руководство" for the document's actual
# category and to prefix it onto the exported filename. Documents with no
# recognized category (manual/imported/no-template) keep both as-is.
_BUILD_TEMPLATE_LABELS = {
    "faq": "FAQ",
    "release-notes": "Release notes",
    "use-case": "Use Case",
    "technical-spec": "Техническая спецификация",
    "api-docs": "Документация API",
}
_DEFAULT_FOOTER_LABEL = "Руководство"

# Fixed conventions shared with markdown.py/html.py: gap nodes never carry
# real text in production (assembled from missing-source detection), and the
# image placeholder text mirrors the HTML template's wording exactly so the
# three formats read consistently.
_GAP_MESSAGE = "Нет данных в источниках"
_IMAGE_PLACEHOLDER_MESSAGE = "Изображение или схема"

# Style names are the real Colvir corporate styles inventoried from
# colvir_v3.dotx's styles.xml (see task-4-report.md for the full mapping
# table and the rationale for each choice).
_HEADING_STYLE_NAMES = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 4",
    5: "Heading 5",
    6: "Heading 6",
}
_TITLE_STYLE = "Colvir_Обложка_Название"
_PARAGRAPH_STYLE = "Colvir_Абзац"
_SECTION_STYLE = "Colvir_Подзаголовок"
_LIST_STYLE = "Colvir_Стиль_М1"
_GAP_STYLE = "Colvir_Внимание"
_CAPTION_STYLE = "Colvir_Рисунок_Подпись"
_TABLE_STYLE = "Colvir_сетка_таблицы"
_TABLE_HEADER_CELL_STYLE = "Colvir_Таблица_заголовок"
_TABLE_BODY_CELL_STYLE = "Colvir_Таблица_текст"

# The template only ships corporate-styled entries for the first three
# outline levels (see task-4-report.md); deeper headings fall back to the
# level-3 look rather than an unstyled default.
_TOC_ENTRY_STYLE_NAMES = {1: "toc 1", 2: "toc 2", 3: "toc 3"}
_TOC_FIELD_SWITCH = ' TOC \\o "1-6" \\h \\z \\u '
_TOC_EMPTY_MESSAGE = "Список разделов пуст"

_LIST_INDENT_LEFT_TWIPS = 720
_LIST_INDENT_HANGING_TWIPS = 360
_FAQ_ITEM_PATTERN = re.compile(
    r"^\s*Вопрос\s*:\s*(?P<question>.+?)\s+Ответ\s*:\s*(?P<answer>.+?)\s*$",
    re.DOTALL,
)


class DocxExporter:
    """Exports WorkingDocuments to styled Word (.docx) files.

    Loads a fresh in-memory copy of the template's base `.docx` asset for
    every render and appends the document's nodes to it using the Colvir
    corporate paragraph/table/list styles. The template's existing
    header/footer, numbering, and styles are preserved untouched; only the
    body's sample/placeholder content is cleared before rendering.
    """

    def __init__(
        self,
        image_loader: ImageLoader | None = None,
        templates_dir: Path | None = None,
    ) -> None:
        """Create a DocxExporter.

        Args:
            image_loader: Resolves a node's image src to bytes + MIME type.
                When None (the default), all images render as placeholders.
            templates_dir: Directory containing the `.docx` assets named by
                a FormattingTemplate's `assets` list. Defaults to the
                built-in formatting/templates catalog directory.
        """
        self._image_loader = image_loader
        self._templates_dir = (templates_dir or _TEMPLATES_DIR).resolve()

    def render(
        self, document: WorkingDocument, template: FormattingTemplate
    ) -> RenderedFile:
        """Render a document to a styled DOCX file.

        Args:
            document: The WorkingDocument to render.
            template: The FormattingTemplate naming the `.docx` base asset
                to load.

        Returns:
            RenderedFile with the DOCX content.
        """
        asset_name = self._asset_named(template, ".docx")
        base_bytes = self._read_asset_bytes(asset_name)

        docx_document = docx.Document(BytesIO(base_bytes))
        self._strip_contextual_spacing(docx_document)
        self._prepare_template_body(docx_document, _cover_title(document.title))
        self._prepare_headers(docx_document, _cover_title(document.title))
        category_label = document_category_label(document)
        if category_label is not None:
            self._prepare_footer(docx_document, category_label)
        docx_document.core_properties.title = document.title

        # Use Case documents render their section headers as fixed form
        # labels (see `_render_use_case_form`), not from these heading
        # nodes, so there is no single rendered paragraph to bookmark --
        # those keep the old flat, non-navigable contents listing.
        heading_nodes = (
            []
            if document.build_template_id == "use-case"
            else list(self._iter_heading_nodes(document.nodes))
        )
        toc_bookmarks = {
            id(node): f"_Toc_docgen_{index}" for index, node in enumerate(heading_nodes)
        }
        self._render_contents(docx_document, document, heading_nodes, toc_bookmarks)
        docx_document.add_page_break()

        list_num_ids: dict[bool, int] = {}
        if document.build_template_id == "use-case":
            self._render_use_case_form(docx_document, document, list_num_ids)
        else:
            for node in document.nodes:
                self._render_node(docx_document, node, list_num_ids, toc_bookmarks)

        buffer = BytesIO()
        docx_document.save(buffer)

        return RenderedFile(
            filename=make_safe_filename(
                filename_title(document), ".docx", reserved_suffix=f"-{template.id}"
            ),
            media_type=_MEDIA_TYPE,
            content=buffer.getvalue(),
        )

    # --- asset loading -----------------------------------------------

    def _asset_named(self, template: FormattingTemplate, suffix: str) -> str:
        """Find the first declared asset ending with `suffix`.

        Only assets listed on the template are ever loaded -- this exporter
        never reads an arbitrary path, only ones the catalog already
        validated for this template.
        """
        matches = [asset for asset in template.assets if asset.endswith(suffix)]
        if not matches:
            raise ValueError(
                f"Шаблон {template.id} не содержит ассет с суффиксом {suffix}"
            )
        return matches[0]

    def _read_asset_bytes(self, name: str) -> bytes:
        """Read a catalog-declared asset file, enforcing containment."""
        path = (self._templates_dir / name).resolve()
        if not path.is_relative_to(self._templates_dir):
            raise ValueError("Недопустимый путь ассета")
        return path.read_bytes()

    # --- body setup ----------------------------------------------------

    def _strip_contextual_spacing(self, docx_document: docx.document.Document) -> None:
        """Remove `w:contextualSpacing` from every paragraph style.

        The corporate template sets this on its body-text style (and a few
        others) -- Word's "no extra space between paragraphs of this style"
        option. Since almost all exported paragraphs share that one style
        (`Colvir_Абзац`), it silently zeroed the visible gap between them
        throughout the whole document even though the style's own
        before/after spacing is non-zero.
        """
        for style in docx_document.styles:
            p_pr = style.element.find(qn("w:pPr"))
            if p_pr is None:
                continue
            contextual_spacing = p_pr.find(qn("w:contextualSpacing"))
            if contextual_spacing is not None:
                p_pr.remove(contextual_spacing)

    def _prepare_template_body(
        self, docx_document: docx.document.Document, title: str
    ) -> None:
        """Retain the corporate cover and remove only template sample content.

        The Colvir template is a designed document, not a blank style
        container. Its first cover paragraphs establish the visual identity;
        the remainder is demonstration text and a sample table of contents.
        Keeping the cover lets the exported DOCX preserve the supplied layout,
        page setup, header and footer rather than recreating an approximation.
        """
        paragraphs = list(docx_document.paragraphs)
        title_index = next(
            (
                index
                for index, paragraph in enumerate(paragraphs)
                if paragraph.style.name == _TITLE_STYLE
            ),
            None,
        )
        if title_index is None:
            raise ValueError("В шаблоне Colvir не найден стиль названия обложки")

        # The supplied template's cover consists of two title lines, the
        # title itself, a subtitle and an introductory paragraph.
        cover_paragraphs = paragraphs[: title_index + 3]
        cover_elements = set()
        for paragraph in cover_paragraphs:
            cover_elements.add(paragraph._p)
            paragraph.text = title if paragraph.style.name == _TITLE_STYLE else ""

        body = docx_document.element.body
        sect_pr = body.find(qn("w:sectPr"))
        for child in list(body):
            if child is not sect_pr and child not in cover_elements:
                body.remove(child)

    def _prepare_headers(
        self, docx_document: docx.document.Document, title: str
    ) -> None:
        """Keep the template header inside the printable area for PDF export."""
        for section in docx_document.sections:
            for paragraph in section.header.paragraphs:
                if "Colvir Banking System" not in paragraph.text:
                    continue
                paragraph.paragraph_format.tab_stops.clear_all()
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    section.page_width - section.left_margin - section.right_margin,
                    WD_TAB_ALIGNMENT.RIGHT,
                )
                paragraph.text = f"Colvir Banking System\t{title}"

    def _prepare_footer(
        self, docx_document: docx.document.Document, label: str
    ) -> None:
        """Swap the template's generic footer label for the document's category.

        The corporate footer literally reads "Руководство" (Guide/Manual) --
        fine for a generic document, wrong for e.g. a FAQ export. Only the
        label run is replaced; the page-number field beside it (and its
        own tab/spacing) is left untouched.
        """
        for section in docx_document.sections:
            for paragraph in section.footer.paragraphs:
                if _DEFAULT_FOOTER_LABEL not in paragraph.text:
                    continue
                for run in paragraph.runs:
                    if run.text == _DEFAULT_FOOTER_LABEL:
                        run.text = label

    def _render_contents(
        self,
        docx_document: docx.document.Document,
        document: WorkingDocument,
        heading_nodes: list[DocumentNode],
        toc_bookmarks: dict[int, str],
    ) -> None:
        """Fill the template's dedicated post-cover page with a real TOC field.

        This is a genuine Word ``{ TOC }`` field: right-click -> "Update
        Field" (or Ctrl+A, F9) recomputes real page numbers and hyperlinks
        from the actual layout. It does *not* auto-update on open -- an
        earlier version set `w:updateFields` for that, but letting Word
        silently recompute the field on its own turned out to be the
        opposite of reliable, so this only ever shows the cached content
        below unless the reader asks for an update. That cached content
        (between the field's ``separate`` and ``end``) is pre-populated from
        the same document nodes as the rest of the export and hyperlinked to
        bookmarks placed around each rendered heading (`_wrap_bookmark`), so
        it reads correctly by default everywhere, including the PDF
        pipeline, which converts through headless LibreOffice and would
        otherwise show an empty/stale TOC.
        """
        title = docx_document.add_paragraph("Оглавление", style=_SECTION_STYLE)
        title.paragraph_format.page_break_before = False

        if document.build_template_id == "use-case":
            for heading, level in self._iter_headings(document.nodes):
                entry = docx_document.add_paragraph(heading, style=_PARAGRAPH_STYLE)
                entry.paragraph_format.left_indent = Cm(0.6 * max(0, level - 1))
            return

        entries = [
            (node.text or "", self._heading_level(node), toc_bookmarks[id(node)])
            for node in heading_nodes
        ]
        self._render_toc_field(docx_document, entries)

    def _render_toc_field(
        self,
        docx_document: docx.document.Document,
        entries: list[tuple[str, int, str]],
    ) -> None:
        if not entries:
            paragraph = docx_document.add_paragraph(
                _TOC_EMPTY_MESSAGE, style=_TOC_ENTRY_STYLE_NAMES[1]
            )
            self._prepend_field_begin(paragraph, _TOC_FIELD_SWITCH)
            self._append_field_end(paragraph)
            return

        section = docx_document.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        hyperlink_style_id = docx_document.styles["Hyperlink"].style_id

        paragraphs = []
        for text, level, bookmark_name in entries:
            style_name = _TOC_ENTRY_STYLE_NAMES[min(max(level, 1), 3)]
            paragraph = docx_document.add_paragraph(style=style_name)
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                usable_width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
            )
            entry_xml = (
                f'<w:hyperlink {nsdecls("w")} w:anchor="{bookmark_name}">'
                f'<w:r><w:rPr><w:rStyle w:val="{hyperlink_style_id}"/></w:rPr>'
                f'<w:t xml:space="preserve">{xml_escape(text)}</w:t></w:r>'
                "<w:r><w:tab/></w:r>"
                '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
                "<w:r><w:instrText xml:space=\"preserve\"> PAGEREF "
                f'{bookmark_name} \\h </w:instrText></w:r>'
                '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
                "<w:r><w:t>1</w:t></w:r>"
                '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
                "</w:hyperlink>"
            )
            paragraph._p.append(parse_xml(entry_xml))
            paragraphs.append(paragraph)

        self._prepend_field_begin(paragraphs[0], _TOC_FIELD_SWITCH)
        self._append_field_end(paragraphs[-1])

    def _prepend_field_begin(self, paragraph: Any, instruction: str) -> None:
        """Insert `begin`/`instrText`/`separate` right after the paragraph's pPr."""
        p_pr = paragraph._p.get_or_add_pPr()
        begin = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="begin"/></w:r>')
        instr = parse_xml(
            f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve">'
            f"{xml_escape(instruction)}</w:instrText></w:r>"
        )
        separate = parse_xml(
            f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="separate"/></w:r>'
        )
        p_pr.addnext(separate)
        p_pr.addnext(instr)
        p_pr.addnext(begin)

    def _append_field_end(self, paragraph: Any) -> None:
        paragraph._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>'))

    def _wrap_bookmark(self, paragraph: Any, name: str) -> None:
        """Surround a rendered heading paragraph with a named bookmark.

        `_render_toc_field` hyperlinks/PAGEREFs the contents entry to this
        same name, so Word's "Update Field" can resolve real page numbers
        and navigate here.
        """
        bookmark_id = name.rsplit("_", 1)[-1]
        start = parse_xml(
            f'<w:bookmarkStart {nsdecls("w")} w:id="{bookmark_id}" w:name="{name}"/>'
        )
        end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bookmark_id}"/>')
        paragraph._p.get_or_add_pPr().addnext(start)
        paragraph._p.append(end)

    def _iter_headings(
        self, nodes: list[DocumentNode]
    ) -> Iterator[tuple[str, int]]:
        for node in self._iter_heading_nodes(nodes):
            yield node.text, self._heading_level(node)

    def _iter_heading_nodes(
        self, nodes: list[DocumentNode]
    ) -> Iterator[DocumentNode]:
        for node in nodes:
            if node.kind is NodeKind.HEADING and node.text:
                yield node
            yield from self._iter_heading_nodes(node.children)

    @staticmethod
    def _heading_level(node: DocumentNode) -> int:
        level = node.data.get("level", 1)
        return level if isinstance(level, int) else 1

    # --- node dispatch ---------------------------------------------------

    def _render_use_case_form(
        self,
        docx_document: docx.document.Document,
        document: WorkingDocument,
        list_num_ids: dict[bool, int],
    ) -> None:
        self._validate_use_case_form(document)
        sections = {
            node.section_id: node
            for node in document.nodes
            if node.kind is NodeKind.HEADING and node.section_id
        }
        self._add_form_heading(docx_document, "Общие сведения")
        metadata = docx_document.add_table(
            rows=len(USE_CASE_METADATA_FIELDS),
            cols=2,
        )
        metadata.style = _TABLE_STYLE
        for row_index, label in enumerate(USE_CASE_METADATA_FIELDS):
            self._set_cell_text(
                metadata.cell(row_index, 0),
                label,
                _TABLE_HEADER_CELL_STYLE,
            )
            self._set_cell_text(
                metadata.cell(row_index, 1),
                "",
                _TABLE_BODY_CELL_STYLE,
            )

        self._add_form_heading(docx_document, "Краткое описание")
        self._render_section_value_paragraph(
            docx_document,
            sections.get("overview"),
        )
        self._add_form_heading(docx_document, "Диаграмма деятельности")
        docx_document.add_paragraph("", style=_PARAGRAPH_STYLE)
        self._add_form_heading(docx_document, "Описание")

        description = docx_document.add_table(
            rows=len(USE_CASE_DESCRIPTION_ROWS),
            cols=3,
        )
        description.style = _TABLE_STYLE
        for row_index, (label, section_id) in enumerate(USE_CASE_DESCRIPTION_ROWS):
            self._set_cell_text(
                description.cell(row_index, 0),
                label,
                _TABLE_HEADER_CELL_STYLE,
            )
            node = sections.get(section_id) if section_id else None
            self._set_use_case_value_cell(
                docx_document,
                description.cell(row_index, 1),
                node,
                list_num_ids,
            )
            self._set_cell_text(
                description.cell(row_index, 2),
                "",
                _TABLE_BODY_CELL_STYLE,
            )

        for heading in (
            "Прототип интерфейса",
            "Диаграмма последовательности",
            "Тест-кейсы",
        ):
            self._add_form_heading(docx_document, heading)
            docx_document.add_paragraph("", style=_PARAGRAPH_STYLE)

        self._add_form_heading(docx_document, "Приложения: Технические спецификации")
        self._add_empty_form_table(
            docx_document,
            USE_CASE_TECHNICAL_HEADERS,
        )
        self._add_form_heading(docx_document, "Ссылки")
        self._add_empty_form_table(docx_document, USE_CASE_LINK_HEADERS)
        self._add_form_heading(docx_document, "История изменений")
        self._add_empty_form_table(
            docx_document,
            USE_CASE_HISTORY_HEADERS,
            empty_rows=2,
        )
        self._add_form_heading(docx_document, "Термины и определения")
        self._render_section_value_paragraph(docx_document, sections.get("terms"))

    @staticmethod
    def _validate_use_case_form(document: WorkingDocument) -> None:
        sections: dict[str, DocumentNode] = {}
        for node in document.nodes:
            if (
                node.kind is not NodeKind.HEADING
                or not node.section_id
                or not (node.text or "").strip()
                or not node.children
            ):
                raise ValueError(
                    "Use Case должен состоять из подписанных структурных разделов"
                )
            if node.section_id in sections:
                raise ValueError(f"Дублируется раздел Use Case: {node.section_id}")
            sections[node.section_id] = node

        missing = {"preconditions", "main-flow", "result"} - sections.keys()
        if missing:
            raise ValueError(
                "Use Case не содержит обязательные структурные разделы: "
                + ", ".join(sorted(missing))
            )

        for node in sections.values():
            for content in node.children:
                if (
                    content.kind is NodeKind.PARAGRAPH
                    and (content.text or "").strip().casefold()
                    == "нет данных в источниках"
                ):
                    raise ValueError(
                        f"Раздел «{node.text}» содержит неструктурный gap"
                    )

        main_flow = sections["main-flow"].children
        if len(main_flow) != 1 or main_flow[0].kind not in {
            NodeKind.LIST,
            NodeKind.GAP,
        }:
            raise ValueError(
                "Основной поток должен быть нумерованным списком или gap"
            )
        if main_flow[0].kind is NodeKind.LIST:
            items = main_flow[0].data.get("items")
            if (
                not main_flow[0].data.get("ordered")
                or not isinstance(items, list)
                or not items
                or any(not isinstance(item, str) or not item.strip() for item in items)
            ):
                raise ValueError(
                    "Основной поток должен содержать отдельные нумерованные шаги"
                )

    def _add_form_heading(
        self,
        docx_document: docx.document.Document,
        text: str,
    ) -> None:
        paragraph = docx_document.add_paragraph(text, style=_SECTION_STYLE)
        paragraph.paragraph_format.page_break_before = False

    def _render_section_value_paragraph(
        self,
        docx_document: docx.document.Document,
        section: DocumentNode | None,
    ) -> None:
        values = self._use_case_values(section)
        docx_document.add_paragraph("\n".join(values), style=_PARAGRAPH_STYLE)

    def _set_use_case_value_cell(
        self,
        docx_document: docx.document.Document,
        cell: Any,
        section: DocumentNode | None,
        list_num_ids: dict[bool, int],
    ) -> None:
        values = self._use_case_values(section)
        paragraph = cell.paragraphs[0]
        paragraph.style = _TABLE_BODY_CELL_STYLE
        if not values:
            paragraph.text = ""
            return
        is_list = bool(
            section
            and any(child.kind is NodeKind.LIST for child in section.children)
        )
        num_id = (
            self._ensure_list_num_id(docx_document, ordered=True, cache=list_num_ids)
            if is_list
            else None
        )
        for index, value in enumerate(values):
            current = paragraph if index == 0 else cell.add_paragraph()
            current.style = _TABLE_BODY_CELL_STYLE
            current.text = value
            if num_id is not None:
                self._apply_list_numbering(current, num_id)

    @staticmethod
    def _use_case_values(section: DocumentNode | None) -> list[str]:
        if section is None or not section.children:
            return []
        values: list[str] = []
        for content in section.children:
            if content.kind is NodeKind.GAP:
                continue
            if content.kind is NodeKind.LIST:
                items = content.data.get("items", [])
                if isinstance(items, list):
                    values.extend(str(item) for item in items)
            elif content.text:
                values.append(content.text)
        return values

    def _add_empty_form_table(
        self,
        docx_document: docx.document.Document,
        headers: tuple[str, ...],
        *,
        empty_rows: int = 1,
    ) -> None:
        table = docx_document.add_table(rows=1 + empty_rows, cols=len(headers))
        table.style = _TABLE_STYLE
        self._mark_header_row(table.rows[0])
        for column, header in enumerate(headers):
            self._set_cell_text(
                table.cell(0, column),
                header,
                _TABLE_HEADER_CELL_STYLE,
            )
            for row in range(1, 1 + empty_rows):
                self._set_cell_text(
                    table.cell(row, column),
                    "",
                    _TABLE_BODY_CELL_STYLE,
                )

    def _render_node(
        self,
        docx_document: docx.document.Document,
        node: DocumentNode,
        list_num_ids: dict[bool, int],
        toc_bookmarks: dict[int, str],
    ) -> None:
        """Render a single node and then all of its children.

        Word documents are a flat sequence of block-level elements, so
        (unlike the HTML exporter's nested DOM) children are appended in
        document order immediately after their parent's own content --
        every node kind may carry children, and all must render.
        """
        if node.kind == NodeKind.HEADING:
            self._render_heading(docx_document, node, toc_bookmarks)
        elif node.kind == NodeKind.PARAGRAPH:
            self._render_paragraph(docx_document, node)
        elif node.kind == NodeKind.LIST:
            self._render_list(docx_document, node, list_num_ids)
        elif node.kind == NodeKind.TABLE:
            self._render_table(docx_document, node)
        elif node.kind == NodeKind.IMAGE:
            self._render_image(docx_document, node)
        elif node.kind == NodeKind.GAP:
            self._render_gap(docx_document, node)

        for child in node.children:
            self._render_node(docx_document, child, list_num_ids, toc_bookmarks)

    def _render_heading(
        self,
        docx_document: docx.document.Document,
        node: DocumentNode,
        toc_bookmarks: dict[int, str],
    ) -> None:
        level = node.data.get("level", 1)
        if not isinstance(level, int):
            level = 1
        level = max(1, min(6, level))
        # FAQ sections are authored as h2 blocks in the editor. The template's
        # Heading 2 carries outline numbering, while its corporate subtitle is
        # the blue, unnumbered visual used by the supplied FAQ layout.
        style = _SECTION_STYLE if level == 2 else _HEADING_STYLE_NAMES[level]
        paragraph = docx_document.add_paragraph(node.text or "", style=style)
        if level == 2:
            # The cover uses this style as a page-level block. FAQ sections
            # need the same typography without starting a blank page first.
            paragraph.paragraph_format.page_break_before = False
        bookmark_name = toc_bookmarks.get(id(node))
        if bookmark_name is not None:
            self._wrap_bookmark(paragraph, bookmark_name)

    def _render_paragraph(
        self, docx_document: docx.document.Document, node: DocumentNode
    ) -> None:
        paragraph = docx_document.add_paragraph(node.text or "", style=_PARAGRAPH_STYLE)
        self._apply_node_style(paragraph, node)

    def _render_list(
        self,
        docx_document: docx.document.Document,
        node: DocumentNode,
        list_num_ids: dict[bool, int],
    ) -> None:
        items = node.data.get("items", [])
        if not isinstance(items, list):
            items = []
        if not items:
            return
        if node.text:
            section = docx_document.add_paragraph(node.text, style=_SECTION_STYLE)
            self._apply_node_style(section, node)
        ordered = bool(node.data.get("ordered", False))
        num_id = self._ensure_list_num_id(docx_document, ordered, list_num_ids)
        item_styles = node.data.get("item_styles")
        for index, item in enumerate(items):
            text = item if isinstance(item, str) else str(item)
            faq_item = _FAQ_ITEM_PATTERN.match(text)
            if faq_item is not None:
                self._render_faq_item(docx_document, faq_item)
                continue
            paragraph = docx_document.add_paragraph(text, style=_LIST_STYLE)
            self._apply_node_style(paragraph, node)
            if isinstance(item_styles, list) and index < len(item_styles):
                self._apply_style_attribute(paragraph, item_styles[index])
            self._apply_list_numbering(paragraph, num_id)

    def _render_faq_item(
        self, docx_document: docx.document.Document, match: re.Match[str]
    ) -> None:
        for label, value in (("Вопрос", match.group("question")), ("Ответ", match.group("answer"))):
            paragraph = docx_document.add_paragraph(style=_PARAGRAPH_STYLE)
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(value)
            paragraph.paragraph_format.space_after = Pt(8 if label == "Вопрос" else 28)

    def _render_table(
        self, docx_document: docx.document.Document, node: DocumentNode
    ) -> None:
        headers = node.data.get("headers", [])
        rows = node.data.get("rows", [])
        if not isinstance(headers, list):
            headers = []
        if not isinstance(rows, list):
            rows = []

        # Only skip rendering if both headers and rows are empty, matching
        # the markdown/html exporters' convention.
        if not headers and not rows:
            return

        if headers:
            col_count = len(headers)
        elif rows and isinstance(rows[0], list) and rows[0]:
            col_count = len(rows[0])
        else:
            col_count = 1

        total_rows = (1 if headers else 0) + len(rows)
        table = docx_document.add_table(rows=total_rows, cols=col_count)
        table.style = _TABLE_STYLE

        row_index = 0
        if headers:
            self._mark_header_row(table.rows[0])
            for col_index in range(col_count):
                text = headers[col_index] if col_index < len(headers) else ""
                self._set_cell_text(
                    table.cell(0, col_index), text, _TABLE_HEADER_CELL_STYLE
                )
            row_index = 1

        for row in rows:
            row_values = row if isinstance(row, list) else []
            padded = list(row_values) + [""] * (col_count - len(row_values))
            padded = padded[:col_count]
            for col_index, value in enumerate(padded):
                self._set_cell_text(
                    table.cell(row_index, col_index), value, _TABLE_BODY_CELL_STYLE
                )
            row_index += 1

    def _render_image(
        self, docx_document: docx.document.Document, node: DocumentNode
    ) -> None:
        src = node.data.get("src")
        alt_text = node.data.get("alt") or node.text or ""
        asset = self._resolve_image(src)

        if asset is not None:
            content, _media_type = asset
            section = docx_document.sections[0]
            usable_width = (
                section.page_width - section.left_margin - section.right_margin
            )
            paragraph = docx_document.add_paragraph()
            run = paragraph.add_run()
            shape = run.add_picture(BytesIO(content), width=usable_width)
            self._set_alt_text(shape, alt_text)
        else:
            docx_document.add_paragraph(_IMAGE_PLACEHOLDER_MESSAGE, style=_PARAGRAPH_STYLE)

        if node.text:
            docx_document.add_paragraph(node.text, style=_CAPTION_STYLE)

    def _render_gap(
        self, docx_document: docx.document.Document, node: DocumentNode
    ) -> None:
        """Render a gap node using the fixed message, never `node.text`.

        Production gap nodes created during assembly carry no text at all
        -- following the same convention markdown.py/html.py already
        established, the message is always the fixed
        "Нет данных в источниках", never node.text.
        """
        docx_document.add_paragraph(_GAP_MESSAGE, style=_GAP_STYLE)

    def _apply_node_style(self, paragraph: Any, node: DocumentNode) -> None:
        self._apply_style_mapping(paragraph, node.data.get("style"))

    def _apply_style_attribute(self, paragraph: Any, value: object) -> None:
        if not isinstance(value, str):
            return
        style = {
            property_name.strip(): property_value.strip()
            for declaration in value.split(";")
            if ":" in declaration
            for property_name, property_value in [declaration.split(":", 1)]
        }
        self._apply_style_mapping(paragraph, style)

    def _apply_style_mapping(self, paragraph: Any, value: object) -> None:
        if not isinstance(value, dict):
            return
        style = value
        if style.get("font-weight") in {"700", "bold"}:
            for run in paragraph.runs:
                run.bold = True
        if style.get("font-style") == "italic":
            for run in paragraph.runs:
                run.italic = True
        if style.get("text-decoration") == "underline":
            for run in paragraph.runs:
                run.underline = True
        if style.get("text-align") in {"left", "center", "right", "justify"}:
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            paragraph.alignment = getattr(
                WD_ALIGN_PARAGRAPH, str(style["text-align"]).upper()
            )

    # --- images ------------------------------------------------------

    def _resolve_image(self, src: object) -> ImageAsset | None:
        """Resolve a node's image src to bytes + MIME type, or None.

        Only locally-storage-resolvable references are embedded. A missing,
        unresolvable, or non-image reference returns None so the caller
        renders the standard placeholder instead of fabricating content or
        fetching an external resource.
        """
        if not isinstance(src, str) or not src or self._image_loader is None:
            return None
        asset = self._image_loader(src)
        if asset is None:
            return None
        _content, media_type = asset
        if not media_type or not media_type.startswith("image/"):
            return None
        return asset

    def _set_alt_text(self, shape: Any, alt_text: str) -> None:
        """Set the drawing's alt text (accessible description) if any."""
        if not alt_text:
            return
        doc_pr = shape._inline.find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", alt_text)

    # --- tables --------------------------------------------------------

    def _mark_header_row(self, row: Any) -> None:
        """Mark a table row to repeat as the header row on each page."""
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    def _set_cell_text(self, cell: Any, value: object, style_name: str) -> None:
        text = value if isinstance(value, str) else str(value)
        paragraph = cell.paragraphs[0]
        paragraph.style = style_name
        paragraph.add_run(text)

    # --- lists / numbering -----------------------------------------------

    def _ensure_list_num_id(
        self,
        docx_document: docx.document.Document,
        ordered: bool,
        cache: dict[bool, int],
    ) -> int:
        """Return a numId for bulleted/numbered list items, creating one if needed.

        The Colvir template ships no generic bullet/numbered *paragraph*
        list style: its one custom "numbering" style (Colvir_Стиль1) is
        wired to the same multilevel numbering used by the heading styles
        (large bold blue numerals), and its Н1/Н2/Н3 styles are unused by
        the template's own sample content (see task-4-report.md). Rather
        than repurpose either of those or inject python-docx's English-
        named "List Bullet"/"List Number" styles into a Cyrillic-branded
        template, this creates a dedicated bullet/decimal numbering
        definition at render time and applies it directly to
        Colvir_Стиль_М1 paragraphs via `w:numPr` -- real Word-native list
        numbering, scoped to this render's in-memory document only.
        """
        if ordered in cache:
            return cache[ordered]

        numbering_element = docx_document.part.numbering_part.element
        existing_abstract_ids = [
            int(element.get(qn("w:abstractNumId")))
            for element in numbering_element.findall(qn("w:abstractNum"))
        ]
        new_abstract_id = max(existing_abstract_ids, default=-1) + 1

        num_fmt = "decimal" if ordered else "bullet"
        lvl_text = "%1." if ordered else "•"
        abstract_xml = (
            f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{new_abstract_id}">'
            "<w:multiLevelType w:val=\"singleLevel\"/>"
            '<w:lvl w:ilvl="0">'
            '<w:start w:val="1"/>'
            f'<w:numFmt w:val="{num_fmt}"/>'
            f'<w:lvlText w:val="{lvl_text}"/>'
            '<w:lvlJc w:val="left"/>'
            "<w:pPr>"
            f'<w:ind w:left="{_LIST_INDENT_LEFT_TWIPS}" w:hanging="{_LIST_INDENT_HANGING_TWIPS}"/>'
            "</w:pPr>"
            "</w:lvl>"
            "</w:abstractNum>"
        )
        abstract_num_element = parse_xml(abstract_xml)

        first_num_element = numbering_element.find(qn("w:num"))
        if first_num_element is not None:
            first_num_element.addprevious(abstract_num_element)
        else:
            numbering_element.append(abstract_num_element)

        num_element = numbering_element.add_num(new_abstract_id)
        num_id = num_element.numId
        cache[ordered] = num_id
        return num_id

    def _apply_list_numbering(self, paragraph: Any, num_id: int, ilvl: int = 0) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl_element = OxmlElement("w:ilvl")
        ilvl_element.set(qn("w:val"), str(ilvl))
        num_id_element = OxmlElement("w:numId")
        num_id_element.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl_element)
        num_pr.append(num_id_element)
        p_pr.append(num_pr)


def document_category_label(document: WorkingDocument) -> str | None:
    """The corporate label for `document`'s build template, or None.

    None means the document has no recognized category (manual/imported/
    no-template) -- callers should leave the template's own generic
    "Руководство" footer and the plain title-derived filename alone.
    """
    if document.build_template_id is None:
        return None
    return _BUILD_TEMPLATE_LABELS.get(document.build_template_id)


def filename_title(document: WorkingDocument) -> str:
    """`document.title`, prefixed with its category label when it has one.

    Shared by `DocxExporter` and `PdfExporter._render_from_docx_template`
    (which converts the same DOCX render) so a FAQ export is recognizable
    as "FAQ-..." in a folder of downloads on both formats, not just DOCX.
    """
    label = document_category_label(document)
    return document.title if label is None else f"{label}-{document.title}"


def _cover_title(title: str) -> str:
    normalized = " ".join(title.split()).strip()
    lowered = normalized.casefold()
    for prefix in ("вопросы и ответы по модулю", "faq по модулю"):
        if lowered.startswith(prefix):
            normalized = f"FAQ {normalized[len(prefix) :].lstrip(' :—-')}"
            break
    if len(normalized) <= 42:
        return normalized or "Документ"
    shortened = normalized[:42].rsplit(" ", 1)[0].rstrip(" ,:;—-")
    return shortened or normalized[:42].rstrip()
