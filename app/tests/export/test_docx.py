import base64
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as OpenDocx
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docgen.documents.schemas import DocumentNode, NodeKind, WorkingDocument
from docgen.export.docx import DocxExporter, local_storage_image_loader
from docgen.export.storage import ExportStorage
from docgen.extraction.docx import DocxExtractor
from docgen.formatting.schemas import FormattingTemplate, OutputFormat
from docgen.models import Source, SourceKind
from docgen.sources.storage import LocalStorage
from docgen.workflows.conversion import conversion_document

# A minimal valid 1x1 transparent PNG, used to exercise the "resolvable src"
# image embedding path without depending on any real asset file.
_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY"
    "42YAAAAASUVORK5CYII="
)


def fake_image_loader(src: str) -> tuple[bytes, str] | None:
    """Test double: resolves a single known src, mirrors the real contract."""
    if src == "images/logo.png":
        return _PNG_BYTES, "image/png"
    return None


@pytest.fixture
def docx_template() -> FormattingTemplate:
    """The real Colvir DOCX template/asset pair, loaded from the catalog dir."""
    return FormattingTemplate(
        id="colvir",
        name="Фирменный стиль Colvir",
        format=OutputFormat.DOCX,
        renderer=OutputFormat.DOCX,
        assets=["colvir.docx"],
    )


def _open(rendered_content: bytes) -> OpenDocx:
    return OpenDocx(BytesIO(rendered_content))


# --- structural / template-fidelity tests ---------------------------------


def test_docx_uses_template_styles_headers_and_tables(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Полный документ",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(kind=NodeKind.HEADING, text="Введение", data={"level": 1}),
            DocumentNode(kind=NodeKind.PARAGRAPH, text="Основной текст."),
            DocumentNode(
                kind=NodeKind.TABLE,
                data={"headers": ["А", "Б"], "rows": [["1", "2"]]},
            ),
        ],
    )

    rendered = DocxExporter(image_loader=fake_image_loader).render(document, docx_template)
    package = _open(rendered.content)

    styles_used = {p.style.name for p in package.paragraphs}
    assert "Heading 1" in styles_used
    assert "Colvir_Абзац" in styles_used
    assert len(package.tables) == 1
    assert package.tables[0].style.name == "Colvir_сетка_таблицы"
    assert package.core_properties.title == document.title


def test_numbered_docx_heading_round_trips_through_editor_as_heading(
    tmp_path: Path,
    docx_template: FormattingTemplate,
) -> None:
    source_path = tmp_path / "numbered-heading.docx"
    source_document = OpenDocx()
    number_id = source_document.styles["List Number"].element.pPr.numPr.numId.val
    heading = source_document.add_heading("Введение", level=2)
    numbering = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(number_id))
    numbering.extend((level, number))
    heading._p.get_or_add_pPr().append(numbering)
    source_document.add_paragraph("Обычный пункт", style="List Number")
    source_document.save(source_path)
    source = Source(
        id="s1",
        project_id="p1",
        kind=SourceKind.FILE,
        display_name=source_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size_bytes=source_path.stat().st_size,
        storage_path="projects/p1/sources/s1.docx",
        status="stored",
    )

    blocks = DocxExtractor().extract_workspace(source, source_path).blocks
    editor_document = conversion_document(
        blocks,
        "Руководство",
        rebase_heading_levels=True,
    )
    rendered = DocxExporter(image_loader=fake_image_loader).render(
        editor_document,
        docx_template,
    )
    exported = _open(rendered.content)

    assert [node.kind for node in editor_document.nodes] == [
        NodeKind.HEADING,
        NodeKind.LIST,
    ]
    exported_heading = next(
        paragraph for paragraph in exported.paragraphs if paragraph.text == "Введение"
    )
    exported_list_item = next(
        paragraph
        for paragraph in exported.paragraphs
        if paragraph.text == "Обычный пункт"
    )
    assert exported_heading.style.name == "Heading 1"
    assert exported_list_item.style.name != "Heading 1"


def test_docx_assembled_use_case_uses_full_corporate_form_with_empty_fields(
    docx_template: FormattingTemplate,
) -> None:
    document = WorkingDocument(
        title="Открытие счёта",
        template_id="use-case",
        nodes=[
            DocumentNode(
                kind=NodeKind.HEADING,
                section_id="main-flow",
                text="Основной поток",
                children=[
                    DocumentNode(
                        kind=NodeKind.LIST,
                        data={
                            "ordered": True,
                            "items": [
                                "Клиент отправляет заявление",
                                "Система открывает счёт",
                            ],
                        },
                    )
                ],
            ),
            DocumentNode(
                kind=NodeKind.HEADING,
                section_id="preconditions",
                text="Предусловия",
                children=[
                    DocumentNode(
                        kind=NodeKind.GAP,
                        flags=["missing-source-data"],
                    )
                ],
            ),
            DocumentNode(
                kind=NodeKind.HEADING,
                section_id="result",
                text="Результат",
                children=[
                    DocumentNode(
                        kind=NodeKind.GAP,
                        flags=["missing-source-data"],
                    )
                ],
            ),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    assert len(package.tables) == 5
    assert package.tables[0].cell(0, 0).text == "Код документа"
    assert package.tables[0].cell(0, 1).text == ""
    assert package.tables[1].cell(0, 0).text == "Область действия"
    assert package.tables[2].cell(0, 0).text == "Наименование"
    assert package.tables[3].cell(0, 0).text == "Ссылка"
    assert package.tables[4].cell(0, 0).text == "Версия документа"
    assert all(table.style.name == "Colvir_сетка_таблицы" for table in package.tables)
    full_text = "\n".join(paragraph.text for paragraph in package.paragraphs)
    assert "Нет данных в источниках" not in full_text
    numbered = [
        paragraph
        for paragraph in package.tables[1].cell(10, 1).paragraphs
        if paragraph.text
    ]
    assert [paragraph.text for paragraph in numbered] == [
        "Клиент отправляет заявление",
        "Система открывает счёт",
    ]
    assert all(
        paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
        for paragraph in numbered
    )


def test_docx_rejects_structurally_invalid_assembled_use_case(
    docx_template: FormattingTemplate,
) -> None:
    document = WorkingDocument(
        title="Склеенный сценарий",
        template_id="use-case",
        nodes=[
            DocumentNode(
                kind=NodeKind.HEADING,
                section_id="preconditions",
                text="Предусловия",
                children=[DocumentNode(kind=NodeKind.GAP)],
            ),
            DocumentNode(
                kind=NodeKind.HEADING,
                section_id="main-flow",
                text="Основной поток",
                children=[
                    DocumentNode(
                        kind=NodeKind.PARAGRAPH,
                        text="1. Клиент отправляет заявку. 2. Система открывает счёт.",
                    )
                ],
            ),
            DocumentNode(
                kind=NodeKind.HEADING,
                section_id="result",
                text="Результат",
                children=[DocumentNode(kind=NodeKind.GAP)],
            ),
        ],
    )

    with pytest.raises(ValueError, match="нумерованным списком"):
        DocxExporter().render(document, docx_template)


def test_docx_reflows_template_header_and_preserves_footer(
    docx_template: FormattingTemplate,
) -> None:
    """The header must fit the printable width in Word and LibreOffice."""
    document = WorkingDocument(title="Документ", template_id="colvir-docx", nodes=[])

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    header_text = package.sections[0].header.paragraphs[0].text
    assert "Colvir Banking System" in header_text
    assert header_text.endswith("\tДокумент")
    tab_stops = package.sections[0].header.paragraphs[0].paragraph_format.tab_stops
    assert len(tab_stops) == 1
    assert tab_stops[0].alignment is WD_TAB_ALIGNMENT.RIGHT
    footer_text = "\n".join(p.text for p in package.sections[0].footer.paragraphs)
    assert "Руководство" in footer_text


def test_docx_replaces_template_sample_body_with_contents(docx_template: FormattingTemplate) -> None:
    """The template sample is removed and its TOC page is regenerated."""
    document = WorkingDocument(title="Документ", template_id="colvir-docx", nodes=[])

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    full_text = "\n".join(p.text for p in package.paragraphs)
    assert "Введите название документа" not in full_text
    assert "Введите заголовок первого уровня" not in full_text
    assert "Оглавление" in full_text
    contents_index, contents = next(
        (index, paragraph)
        for index, paragraph in enumerate(package.paragraphs)
        if paragraph.text == "Оглавление"
    )
    assert contents.paragraph_format.page_break_before is True
    assert contents.paragraph_format.space_before.pt == 0
    assert contents_index == 3


def test_docx_contents_uses_document_headings(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Документ",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(kind=NodeKind.HEADING, text="Общие вопросы", data={"level": 2}),
            DocumentNode(
                kind=NodeKind.HEADING,
                text="Настройки",
                data={"level": 3},
            ),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    body_headings = {
        p.text: p for p in package.paragraphs if p.text in ("Общие вопросы", "Настройки")
    }
    assert body_headings["Общие вопросы"].style.name == "Colvir_Подзаголовок"
    assert body_headings["Настройки"].style.name == "Heading 3"

    toc_entries = {
        p.text.split("\t", 1)[0]: p
        for p in package.paragraphs
        if p.text.startswith("Общие вопросы\t") or p.text.startswith("Настройки\t")
    }
    assert toc_entries["Общие вопросы"].style.name == "toc 2"
    assert toc_entries["Настройки"].style.name == "toc 3"

    def bookmark_name(paragraph) -> str:
        return paragraph._p.find(qn("w:bookmarkStart")).get(qn("w:name"))

    def anchor_name(paragraph) -> str:
        return paragraph._p.find(qn("w:hyperlink")).get(qn("w:anchor"))

    # The contents entry is a real hyperlink/PAGEREF wired to the same
    # bookmark that wraps the actual rendered heading -- not a lookalike.
    assert bookmark_name(body_headings["Общие вопросы"]) == anchor_name(
        toc_entries["Общие вопросы"]
    )
    assert bookmark_name(body_headings["Настройки"]) == anchor_name(
        toc_entries["Настройки"]
    )


def test_docx_contents_includes_titled_list_sections(
    docx_template: FormattingTemplate,
) -> None:
    """Assembled FAQ documents title a section via `.text` on a LIST node,
    not a HEADING node (see assemble.py's FAQ assembly instructions) --
    the contents page must still find and bookmark it."""
    document = WorkingDocument(
        title="Документ",
        template_id="faq",
        nodes=[
            DocumentNode(
                kind=NodeKind.LIST,
                text="Общие вопросы",
                data={"items": ["Вопрос: А?\nОтвет: Б."]},
            ),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    body_section = next(
        p for p in package.paragraphs if p.text == "Общие вопросы"
    )
    assert body_section.style.name == "Colvir_Подзаголовок"

    toc_entry = next(
        p for p in package.paragraphs if p.text.startswith("Общие вопросы\t")
    )
    assert toc_entry.style.name == "toc 2"

    bookmark_name = body_section._p.find(qn("w:bookmarkStart")).get(qn("w:name"))
    anchor_name = toc_entry._p.find(qn("w:hyperlink")).get(qn("w:anchor"))
    assert bookmark_name == anchor_name


def test_docx_contents_skips_untitled_and_empty_lists(
    docx_template: FormattingTemplate,
) -> None:
    document = WorkingDocument(
        title="Документ",
        template_id="faq",
        nodes=[
            DocumentNode(kind=NodeKind.LIST, data={"items": ["Обычный пункт"]}),
            DocumentNode(kind=NodeKind.LIST, text="Пустой раздел", data={"items": []}),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    full_text = "\n".join(p.text for p in package.paragraphs)
    assert "Список разделов пуст" in full_text
    assert "Пустой раздел" not in full_text


def test_docx_contents_is_a_real_updatable_toc_field(
    docx_template: FormattingTemplate,
) -> None:
    document = WorkingDocument(
        title="Документ",
        template_id="colvir-docx",
        nodes=[DocumentNode(kind=NodeKind.HEADING, text="Раздел", data={"level": 1})],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    instr_texts = [
        element.text
        for element in package.element.body.iter(qn("w:instrText"))
    ]
    assert any(text and "TOC" in text for text in instr_texts)
    field_chars = [
        element.get(qn("w:fldCharType"))
        for element in package.element.body.iter(qn("w:fldChar"))
    ]
    assert "begin" in field_chars
    assert "separate" in field_chars
    assert "end" in field_chars

    # Word refreshes the genuine TOC/PAGEREF fields when the file opens.
    settings = package.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    assert update_fields is not None
    assert update_fields.get(qn("w:val")) == "true"


def test_docx_contents_with_no_headings_shows_placeholder(
    docx_template: FormattingTemplate,
) -> None:
    document = WorkingDocument(title="Документ", template_id="colvir-docx", nodes=[])

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    full_text = "\n".join(p.text for p in package.paragraphs)
    assert "Список разделов пуст" in full_text


def test_docx_sets_title_paragraph_and_metadata(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(title="Заголовок пакета", template_id="colvir-docx", nodes=[])

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    assert package.core_properties.title == "Заголовок пакета"
    title_paragraphs = [
        p for p in package.paragraphs if p.style.name == "Colvir_Обложка_Название"
    ]
    assert any(p.text == "Заголовок пакета" for p in title_paragraphs)


def test_docx_uses_short_cover_title_and_faq_paragraphs(
    docx_template: FormattingTemplate,
) -> None:
    document = WorkingDocument(
        title="Вопросы и ответы по модулю «Главная Книга» Colvir Banking System",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(kind=NodeKind.HEADING, text="Общие вопросы", data={"level": 2}),
            DocumentNode(
                kind=NodeKind.LIST,
                data={
                    "items": [
                        (
                            "Вопрос: Для чего предназначен модуль? "
                            "Ответ: Для ведения Главной Книги."
                        )
                    ]
                },
            ),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    title = next(
        p for p in package.paragraphs if p.style.name == "Colvir_Обложка_Название"
    )
    assert title.text == "FAQ «Главная Книга» Colvir Banking System"
    header_text = package.sections[0].header.paragraphs[0].text
    assert header_text.endswith("\tFAQ «Главная Книга» Colvir Banking System")

    heading = next(
        p
        for p in package.paragraphs
        if p.text == "Общие вопросы" and p.style.name == "Colvir_Подзаголовок"
    )
    assert heading.style.name == "Colvir_Подзаголовок"
    assert heading.paragraph_format.page_break_before is False

    question = next(p for p in package.paragraphs if p.text.startswith("Вопрос:"))
    answer = next(p for p in package.paragraphs if p.text.startswith("Ответ:"))
    assert question.style.name == "Colvir_Абзац"
    assert answer.style.name == "Colvir_Абзац"
    assert question.runs[0].bold is True
    assert answer.runs[0].bold is True
    assert question._p.pPr.find(qn("w:numPr")) is None
    assert question.paragraph_format.space_after.twips == Pt(8).twips
    assert answer.paragraph_format.space_after.twips == Pt(28).twips


# --- heading level mapping --------------------------------------------------


@pytest.mark.parametrize(
    "level,expected_style",
    [(1, "Heading 1"), (2, "Colvir_Подзаголовок"), (3, "Heading 3"), (6, "Heading 6")],
)
def test_docx_heading_level_maps_to_style(
    docx_template: FormattingTemplate, level: int, expected_style: str
) -> None:
    document = WorkingDocument(
        title="Документ",
        template_id="colvir-docx",
        nodes=[DocumentNode(kind=NodeKind.HEADING, text="Раздел", data={"level": level})],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    body_heading = next(p for p in package.paragraphs if p.text == "Раздел")
    assert body_heading.style.name == expected_style
    toc_entry = next(p for p in package.paragraphs if p.text.startswith("Раздел\t"))
    assert toc_entry.style.name in ("toc 1", "toc 2", "toc 3")


def test_docx_heading_level_is_clamped(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Документ",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(kind=NodeKind.HEADING, text="Слишком глубоко", data={"level": 99}),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    matching = [p for p in package.paragraphs if p.text == "Слишком глубоко"]
    assert matching[-1].style.name == "Heading 6"


# --- tables: header-less / rows-less edge cases -----------------------------


def test_docx_renders_table_with_headers_no_rows(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Таблица только с заголовками",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(
                kind=NodeKind.TABLE,
                data={"headers": ["Колонка А", "Колонка Б"], "rows": []},
            ),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    assert len(package.tables) == 1
    table = package.tables[0]
    assert len(table.rows) == 1
    assert [c.text for c in table.rows[0].cells] == ["Колонка А", "Колонка Б"]
    assert table.rows[0].cells[0].paragraphs[0].style.name == "Colvir_Таблица_заголовок"


def test_docx_renders_table_without_headers(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Таблица без заголовков",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(
                kind=NodeKind.TABLE,
                data={"rows": [["Значение 1", "Значение 2"]]},
            ),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    table = package.tables[0]
    assert len(table.rows) == 1
    assert [c.text for c in table.rows[0].cells] == ["Значение 1", "Значение 2"]
    assert table.rows[0].cells[0].paragraphs[0].style.name == "Colvir_Таблица_текст"


def test_docx_skips_empty_table(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Пустая таблица",
        template_id="colvir-docx",
        nodes=[DocumentNode(kind=NodeKind.TABLE, data={"rows": []})],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    assert len(package.tables) == 0


def test_docx_pads_ragged_table_rows(docx_template: FormattingTemplate) -> None:
    """Rows shorter/longer than the header count must be padded/truncated."""
    document = WorkingDocument(
        title="Неровная таблица",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(
                kind=NodeKind.TABLE,
                data={"headers": ["А", "Б", "В"], "rows": [["1"], ["1", "2", "3", "4"]]},
            ),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    table = package.tables[0]
    assert len(table.rows) == 3
    assert [c.text for c in table.rows[1].cells] == ["1", "", ""]
    assert [c.text for c in table.rows[2].cells] == ["1", "2", "3"]


# --- gap nodes ---------------------------------------------------------------


def test_docx_renders_textless_gap_node(docx_template: FormattingTemplate) -> None:
    """Production gap nodes created during assembly have no text attribute."""
    document = WorkingDocument(
        title="Документ с пробелом",
        template_id="colvir-docx",
        nodes=[DocumentNode(kind=NodeKind.GAP)],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    gap_paragraphs = [p for p in package.paragraphs if p.style.name == "Colvir_Внимание"]
    assert len(gap_paragraphs) == 1
    assert gap_paragraphs[0].text == "Нет данных в источниках"


def test_docx_gap_node_ignores_its_own_text(docx_template: FormattingTemplate) -> None:
    """Even if a gap node somehow carries text, the fixed message wins."""
    document = WorkingDocument(
        title="Документ",
        template_id="colvir-docx",
        nodes=[DocumentNode(kind=NodeKind.GAP, text="какой-то произвольный текст")],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    full_text = "\n".join(p.text for p in package.paragraphs)
    assert "Нет данных в источниках" in full_text
    assert "какой-то произвольный текст" not in full_text


# --- images: resolution fallback pattern (Task 3 ruling) -------------------


def test_docx_embeds_resolvable_image(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Документ с изображением",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(
                kind=NodeKind.IMAGE,
                data={"src": "images/logo.png", "alt": "Логотип"},
                text="Подпись к рисунку",
            ),
        ],
    )

    rendered = DocxExporter(image_loader=fake_image_loader).render(document, docx_template)
    package = _open(rendered.content)

    assert len(package.inline_shapes) == 1
    caption_paragraphs = [
        p for p in package.paragraphs if p.style.name == "Colvir_Рисунок_Подпись"
    ]
    assert any(p.text == "Подпись к рисунку" for p in caption_paragraphs)


def test_docx_image_without_src_renders_placeholder(docx_template: FormattingTemplate) -> None:
    """Production AI-assembled image nodes carry no data.src at all.

    Must not fabricate or attempt an embed; must render the same
    placeholder wording as the HTML exporter's fallback.
    """
    document = WorkingDocument(
        title="Документ без источника изображения",
        template_id="colvir-docx",
        nodes=[DocumentNode(kind=NodeKind.IMAGE, text="Схема архитектуры")],
    )

    rendered = DocxExporter(image_loader=fake_image_loader).render(document, docx_template)
    package = _open(rendered.content)

    assert len(package.inline_shapes) == 0
    full_text = "\n".join(p.text for p in package.paragraphs)
    assert "Изображение или схема" in full_text
    assert "Схема архитектуры" in full_text


def test_docx_image_with_unresolvable_src_renders_placeholder(
    docx_template: FormattingTemplate,
) -> None:
    """A src the loader cannot resolve (e.g. deleted file) falls back safely."""
    document = WorkingDocument(
        title="Документ с недоступным изображением",
        template_id="colvir-docx",
        nodes=[DocumentNode(kind=NodeKind.IMAGE, data={"src": "images/missing.png"})],
    )

    rendered = DocxExporter(image_loader=fake_image_loader).render(document, docx_template)
    package = _open(rendered.content)

    assert len(package.inline_shapes) == 0
    full_text = "\n".join(p.text for p in package.paragraphs)
    assert "Изображение или схема" in full_text


def test_docx_renders_without_image_loader(docx_template: FormattingTemplate) -> None:
    """DocxExporter must work with no image_loader at all (defaults to placeholders)."""
    document = WorkingDocument(
        title="Документ",
        template_id="colvir-docx",
        nodes=[DocumentNode(kind=NodeKind.IMAGE, data={"src": "images/logo.png"})],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    assert len(package.inline_shapes) == 0


# --- lists -------------------------------------------------------------------


def test_docx_unordered_list_applies_bullet_numbering(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Документ",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(
                kind=NodeKind.LIST, data={"items": ["Первый", "Второй"], "ordered": False}
            ),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    item_paragraphs = [p for p in package.paragraphs if p.text in ("Первый", "Второй")]
    assert len(item_paragraphs) == 2
    for paragraph in item_paragraphs:
        assert paragraph.style.name == "Colvir_Стиль_М1"
        num_pr = paragraph._p.pPr.find(qn("w:numPr"))
        assert num_pr is not None


def test_docx_ordered_and_unordered_lists_use_distinct_numbering(
    docx_template: FormattingTemplate,
) -> None:
    document = WorkingDocument(
        title="Документ",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(kind=NodeKind.LIST, data={"items": ["А"], "ordered": True}),
            DocumentNode(kind=NodeKind.LIST, data={"items": ["Б"], "ordered": False}),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    def num_id_for(text: str) -> str:
        paragraph = next(p for p in package.paragraphs if p.text == text)
        num_id_element = paragraph._p.pPr.find(qn("w:numPr")).find(qn("w:numId"))
        return num_id_element.get(qn("w:val"))

    assert num_id_for("А") != num_id_for("Б")


def test_docx_empty_list_renders_nothing(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Документ",
        template_id="colvir-docx",
        nodes=[DocumentNode(kind=NodeKind.LIST, data={"items": []})],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    # The cover title, the "Оглавление" title, and the empty-contents
    # placeholder (there are no headings to list) are the only text.
    assert len([p for p in package.paragraphs if p.text.strip()]) == 3


# --- nested children on every node kind -------------------------------------


def test_docx_renders_nested_children_for_every_kind(docx_template: FormattingTemplate) -> None:
    """Every node kind may carry children; all must be rendered."""
    document = WorkingDocument(
        title="Документ с вложенными узлами",
        template_id="colvir-docx",
        nodes=[
            DocumentNode(
                kind=NodeKind.HEADING,
                text="Заголовок",
                data={"level": 1},
                children=[
                    DocumentNode(
                        kind=NodeKind.PARAGRAPH,
                        text="Вложенный абзац",
                        children=[
                            DocumentNode(
                                kind=NodeKind.LIST,
                                data={"items": ["А", "Б"]},
                                children=[
                                    DocumentNode(
                                        kind=NodeKind.TABLE,
                                        data={"rows": [["1", "2"]]},
                                        children=[
                                            DocumentNode(
                                                kind=NodeKind.GAP,
                                                children=[
                                                    DocumentNode(
                                                        kind=NodeKind.IMAGE,
                                                        text="Вложенное изображение",
                                                    )
                                                ],
                                            )
                                        ],
                                    )
                                ],
                            )
                        ],
                    )
                ],
            ),
        ],
    )

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    full_text = "\n".join(p.text for p in package.paragraphs)
    assert "Заголовок" in full_text
    assert "Вложенный абзац" in full_text
    assert "А" in full_text and "Б" in full_text
    assert len(package.tables) == 1
    assert package.tables[0].rows[0].cells[0].text == "1"
    assert "Нет данных в источниках" in full_text
    assert "Вложенное изображение" in full_text


# --- filename / media type ---------------------------------------------------


def test_docx_footer_shows_category_label_for_faq(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(title="Общие вопросы", template_id="faq", nodes=[])

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    footer_text = "\n".join(p.text for p in package.sections[0].footer.paragraphs)
    assert "FAQ" in footer_text
    assert "Руководство" not in footer_text


def test_docx_footer_keeps_default_label_without_category(
    docx_template: FormattingTemplate,
) -> None:
    document = WorkingDocument(title="Документ", template_id="colvir-docx", nodes=[])

    rendered = DocxExporter().render(document, docx_template)
    package = _open(rendered.content)

    footer_text = "\n".join(p.text for p in package.sections[0].footer.paragraphs)
    assert "Руководство" in footer_text


def test_docx_filename_includes_category_label(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(title="Общие вопросы", template_id="faq", nodes=[])

    rendered = DocxExporter().render(document, docx_template)

    assert rendered.filename.startswith("FAQ-")


def test_docx_filename_does_not_double_up_a_title_that_already_has_the_label(
    docx_template: FormattingTemplate,
) -> None:
    """assemble.py defaults a FAQ's title to "FAQ по материалам источников" --
    the label must not be prefixed again on top of that."""
    document = WorkingDocument(
        title="FAQ Интеграция с INHOUSE через AMS", template_id="faq", nodes=[]
    )

    rendered = DocxExporter().render(document, docx_template)

    assert rendered.filename.startswith("FAQ-Интеграция")
    assert "FAQ-FAQ" not in rendered.filename


def test_docx_filename_and_media_type(docx_template: FormattingTemplate) -> None:
    document = WorkingDocument(
        title="Мой прекрасный документ", template_id="colvir-docx", nodes=[]
    )

    rendered = DocxExporter().render(document, docx_template)

    assert rendered.filename.endswith(".docx")
    assert (
        rendered.media_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# --- local_storage_image_loader re-export (shared with HtmlExporter) -------


def test_local_storage_image_loader_resolves_stored_image(tmp_path) -> None:
    storage = LocalStorage(tmp_path)
    saved = storage.save("project-1", "source-1", "logo.png", BytesIO(_PNG_BYTES))
    loader = local_storage_image_loader(storage)

    result = loader(saved.relative_path)

    assert result is not None
    content, media_type = result
    assert content == _PNG_BYTES
    assert media_type == "image/png"


# --- filename byte-length safety (finding 6) -------------------------------


def test_docx_long_cyrillic_title_produces_storable_filename(
    docx_template: FormattingTemplate, tmp_path: Path
) -> None:
    """A 150+ character Cyrillic title must never produce a filename that
    overflows the filesystem's 255-byte limit once ExportStorage appends
    `-{template_id}` -- reproduced end-to-end via the real storage layer."""
    long_title = "Очень длинное название регламента для банковского документа " * 4
    assert len(long_title) > 150
    document = WorkingDocument(title=long_title, template_id="colvir", nodes=[])

    rendered = DocxExporter().render(document, docx_template)
    storage = ExportStorage(tmp_path / "data")

    stored = storage.save("proj-1", OutputFormat.DOCX, docx_template.id, rendered)

    assert len(stored.filename.encode("utf-8")) <= 255
    assert storage.resolve(stored.relative_path).is_file()
