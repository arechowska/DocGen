from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from docgen.extraction.docx import DocxExtractor
from docgen.extraction.registry import ExtractionError
from docgen.extraction.schemas import BlockKind
from docgen.models import Source, SourceKind


def make_source() -> Source:
    return Source(
        id="s1",
        project_id="p1",
        kind=SourceKind.FILE,
        display_name="input.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size_bytes=1,
        storage_path="projects/p1/sources/s1.docx",
        status="stored",
    )

def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_numbering(paragraph, num_id: int, level: int) -> None:
    numbering = OxmlElement("w:numPr")
    level_element = OxmlElement("w:ilvl")
    level_element.set(qn("w:val"), str(level))
    number_id_element = OxmlElement("w:numId")
    number_id_element.set(qn("w:val"), str(num_id))
    numbering.extend((level_element, number_id_element))
    paragraph._p.get_or_add_pPr().append(numbering)


@pytest.fixture
def workspace_fidelity_docx(tmp_path: Path) -> Path:
    path = tmp_path / "workspace-fidelity.docx"
    document = Document()
    document.add_heading("Guide", level=2)

    paragraph = document.add_paragraph()
    bold = paragraph.add_run("Bold")
    bold.bold = True
    italic = paragraph.add_run("Italic")
    italic.italic = True
    underline = paragraph.add_run("Under")
    underline.underline = True
    strike = paragraph.add_run("Strike")
    strike.font.strike = True
    red = paragraph.add_run("Red")
    red.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    _add_hyperlink(paragraph, "Link", "https://example.test")

    number_id = document.styles["List Number"].element.pPr.numPr.numId.val
    bullet_id = document.styles["List Bullet"].element.pPr.numPr.numId.val
    first = document.add_paragraph("First")
    _set_numbering(first, number_id, 0)
    second = document.add_paragraph("Second")
    _set_numbering(second, number_id, 0)
    nested = document.add_paragraph("Nested")
    _set_numbering(nested, bullet_id, 1)
    document.save(path)
    return path



@pytest.fixture
def cyrillic_builtin_styles_docx(tmp_path: Path) -> Path:
    path = tmp_path / "cyrillic-builtins.docx"
    document = Document()
    heading_style = document.styles["Heading 2"]
    heading_style.element.xpath("./w:name")[0].set(
        qn("w:val"), "\u0417\u0430\u0433\u043e\u043b\u043e\u0432\u043e\u043a \u0432\u0442\u043e\u0440\u043e\u0433\u043e \u0443\u0440\u043e\u0432\u043d\u044f"
    )
    list_style = document.styles["List Bullet"]
    list_style.element.xpath("./w:name")[0].set(
        qn("w:val"), "\u041c\u0430\u0440\u043a\u0438\u0440\u043e\u0432\u0430\u043d\u043d\u044b\u0439 \u0441\u043f\u0438\u0441\u043e\u043a"
    )
    document.add_paragraph("\u0420\u0430\u0437\u0434\u0435\u043b", style=heading_style)
    document.add_paragraph("\u041f\u0443\u043d\u043a\u0442", style=list_style)
    document.save(path)
    return path


@pytest.fixture
def inherited_semantic_styles_docx(tmp_path: Path) -> Path:
    path = tmp_path / "inherited-styles.docx"
    document = Document()

    outline_style = document.styles.add_style(
        "\u0420\u0430\u0437\u0434\u0435\u043b \u043f\u0440\u043e\u0435\u043a\u0442\u0430", WD_STYLE_TYPE.PARAGRAPH
    )
    outline_level = OxmlElement("w:outlineLvl")
    outline_level.set(qn("w:val"), "3")
    outline_style.element.get_or_add_pPr().append(outline_level)

    numbered_style = document.styles["List Number"]
    numbered_style.element.xpath("./w:name")[0].set(
        qn("w:val"), "\u041d\u0443\u043c\u0435\u0440\u043e\u0432\u0430\u043d\u043d\u044b\u0439 \u0441\u043f\u0438\u0441\u043e\u043a"
    )
    inherited_list_style = document.styles.add_style(
        "\u041f\u0435\u0440\u0435\u0447\u0435\u043d\u044c \u043f\u0440\u043e\u0435\u043a\u0442\u0430", WD_STYLE_TYPE.PARAGRAPH
    )
    inherited_list_style.base_style = numbered_style

    document.add_paragraph("\u0413\u043b\u0443\u0431\u0438\u043d\u0430", style=outline_style)
    document.add_paragraph("\u0428\u0430\u0433", style=inherited_list_style)
    document.save(path)
    return path


def test_docx_preserves_document_order_and_structure(tmp_path: Path) -> None:
    path = tmp_path / "input.docx"
    document = Document()
    document.add_heading("Раздел", level=1)
    document.add_paragraph("Обычный текст")
    document.add_paragraph("Пункт", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Ключ"
    table.cell(0, 1).text = "Значение"
    document.save(path)

    result = DocxExtractor().extract(make_source(), path)

    assert [block.kind for block in result.blocks] == [
        BlockKind.HEADING,
        BlockKind.TEXT,
        BlockKind.LIST,
        BlockKind.TABLE,
    ]
    assert [block.provenance[0].locator for block in result.blocks] == [
        "paragraph:1",
        "paragraph:2",
        "paragraph:3",
        "table:1",
    ]
    assert result.blocks[0].data == {"level": 1}
    assert result.blocks[2].data == {"style": "List Bullet"}
    assert result.blocks[3].data == {"rows": [["Ключ", "Значение"]]}


def test_docx_workspace_extraction_preserves_editor_safe_rich_content_and_lists(
    workspace_fidelity_docx: Path,
) -> None:
    extractor = DocxExtractor()

    regular = extractor.extract(make_source(), workspace_fidelity_docx)
    result = extractor.extract_workspace(make_source(), workspace_fidelity_docx)

    assert [block.kind for block in regular.blocks] == [
        BlockKind.HEADING,
        BlockKind.TEXT,
        BlockKind.LIST,
        BlockKind.LIST,
        BlockKind.LIST,
    ]
    assert result.blocks[1].data["html"] == (
        '<strong>Bold</strong><em>Italic</em><u>Under</u><s>Strike</s>'
        '<span style="color:#ff0000">Red</span><a href="https://example.test">Link</a>'
    )
    assert result.blocks[2].data == {
        "ordered": True,
        "items": ["First", "Second"],
        "items_html": ["First", "Second<ul><li>Nested</li></ul>"],
    }


def test_docx_uses_builtin_style_ids_when_display_names_are_cyrillic(
    cyrillic_builtin_styles_docx: Path,
) -> None:
    result = DocxExtractor().extract(make_source(), cyrillic_builtin_styles_docx)

    assert [block.kind for block in result.blocks] == [BlockKind.HEADING, BlockKind.LIST]
    assert result.blocks[0].data == {"level": 2}


def test_docx_uses_outline_level_for_custom_heading(
    inherited_semantic_styles_docx: Path,
) -> None:
    result = DocxExtractor().extract(make_source(), inherited_semantic_styles_docx)

    assert result.blocks[0].kind == BlockKind.HEADING
    assert result.blocks[0].data == {"level": 4}


def test_docx_uses_numbering_inherited_from_base_style(
    inherited_semantic_styles_docx: Path,
) -> None:
    result = DocxExtractor().extract(make_source(), inherited_semantic_styles_docx)

    assert result.blocks[1].kind == BlockKind.LIST


def test_docx_demotes_long_heading_styled_paragraph_to_text(tmp_path: Path) -> None:
    path = tmp_path / "input.docx"
    document = Document()
    long_text = " ".join(["Слово"] * 25) + " составляют длинный абзац, а не заголовок."
    document.add_paragraph(long_text, style="Heading 2")
    document.save(path)

    result = DocxExtractor().extract(make_source(), path)

    assert result.blocks[0].kind == BlockKind.TEXT
    assert result.blocks[0].text == long_text


def test_docx_demotes_long_outline_level_paragraph_to_text(tmp_path: Path) -> None:
    path = tmp_path / "input.docx"
    document = Document()
    outline_style = document.styles.add_style(
        "Раздел проекта", WD_STYLE_TYPE.PARAGRAPH
    )
    outline_level = OxmlElement("w:outlineLvl")
    outline_level.set(qn("w:val"), "3")
    outline_style.element.get_or_add_pPr().append(outline_level)
    long_text = " ".join(["Слово"] * 25) + " составляют длинный абзац, а не заголовок."
    document.add_paragraph(long_text, style=outline_style)
    document.save(path)

    result = DocxExtractor().extract(make_source(), path)

    assert result.blocks[0].kind == BlockKind.TEXT


def test_docx_extraction_has_repeatable_block_ids(tmp_path: Path) -> None:
    path = tmp_path / "input.docx"
    document = Document()
    document.add_paragraph("Повторяемый текст")
    document.save(path)
    extractor = DocxExtractor()

    first_result = extractor.extract(make_source(), path)
    second_result = extractor.extract(make_source(), path)

    assert [block.id for block in first_result.blocks] == [block.id for block in second_result.blocks]


def test_docx_counts_virtual_page_boundaries(tmp_path: Path) -> None:
    path = tmp_path / "input.docx"
    extractor = DocxExtractor()
    document = Document()
    document.add_paragraph("a" * 1800)
    document.save(path)

    exact_page = extractor.extract(make_source(), path)

    document = Document()
    document.add_paragraph("a" * 1801)
    document.save(path)
    next_page = extractor.extract(make_source(), path)

    assert exact_page.page_units == 1
    assert next_page.page_units == 2


def test_docx_archive_budgets_accept_exact_boundaries(tmp_path: Path) -> None:
    path = tmp_path / "boundary.docx"
    document = Document()
    document.add_paragraph("Документ")
    document.save(path)
    with ZipFile(path) as archive:
        entry_count = len(archive.infolist())
        expanded_bytes = sum(entry.file_size for entry in archive.infolist())

    result = DocxExtractor(
        max_archive_entries=entry_count,
        max_archive_uncompressed_bytes=expanded_bytes,
    ).extract(make_source(), path)

    assert result.blocks[0].text == "Документ"


@pytest.mark.parametrize("budget", ["entries", "expanded"])
def test_docx_archive_budget_rejected_before_ooxml_parse(
    tmp_path: Path, budget: str
) -> None:
    path = tmp_path / "oversized.docx"
    document = Document()
    document.add_paragraph("Документ")
    document.save(path)
    with ZipFile(path) as archive:
        entry_count = len(archive.infolist())
        expanded_bytes = sum(entry.file_size for entry in archive.infolist())
    kwargs = {
        "max_archive_entries": entry_count,
        "max_archive_uncompressed_bytes": expanded_bytes,
    }
    if budget == "entries":
        kwargs["max_archive_entries"] -= 1
    else:
        kwargs["max_archive_uncompressed_bytes"] -= 1

    with pytest.raises(ExtractionError, match="Архив DOCX превышает допустимый объём"):
        DocxExtractor(**kwargs).extract(make_source(), path)


def test_docx_returns_safe_error_when_file_cannot_be_read(tmp_path: Path) -> None:
    with pytest.raises(ExtractionError, match="Не удалось прочитать DOCX-файл"):
        DocxExtractor().extract(make_source(), tmp_path / "missing.docx")


def test_docx_returns_safe_error_for_zip_missing_ooxml_members(tmp_path: Path) -> None:
    path = tmp_path / "missing-members.docx"
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("placeholder.txt", "not an OOXML package")

    with pytest.raises(ExtractionError, match="Не удалось прочитать DOCX-файл"):
        DocxExtractor().extract(make_source(), path)


def test_docx_returns_safe_error_for_malformed_package_xml(tmp_path: Path) -> None:
    path = tmp_path / "malformed-xml.docx"
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types>")
        archive.writestr("_rels/.rels", "<Relationships>")
        archive.writestr("word/document.xml", "<document>")

    with pytest.raises(ExtractionError, match="Не удалось прочитать DOCX-файл"):
        DocxExtractor().extract(make_source(), path)


def test_docx_workspace_counts_nested_list_text_for_page_limits(tmp_path: Path) -> None:
    path = tmp_path / "nested-page-count.docx"
    document = Document()
    number_id = document.styles["List Number"].element.pPr.numPr.numId.val
    bullet_id = document.styles["List Bullet"].element.pPr.numPr.numId.val
    parent = document.add_paragraph("Parent")
    _set_numbering(parent, number_id, 0)
    nested_text = "x" * 1800
    nested = document.add_paragraph(nested_text)
    _set_numbering(nested, bullet_id, 1)
    document.save(path)

    result = DocxExtractor().extract_workspace(make_source(), path)

    assert result.blocks[0].text == f"Parent\n{nested_text}"
    assert result.page_units == 2
