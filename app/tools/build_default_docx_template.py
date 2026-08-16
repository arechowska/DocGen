#!/usr/bin/env python3
"""Build the runtime DOCX base asset from the Colvir corporate template.

The user supplied a real corporate Word template, ``colvir_v3.dotx``. It is a
Word *template* package (OOXML content type
``...wordprocessingml.template.main+xml`` for ``/word/document.xml``), and
python-docx's ``Document()`` opener only accepts the *document* content type
(``...wordprocessingml.document.main+xml``) -- opening a ``.dotx`` directly
raises ``ValueError``.

This script does **not** rebuild styles, numbering, headers, or footers from
scratch. It repackages ``colvir_v3.dotx`` as-is: it opens the source as a
zip, patches the single ``Content_Types`` Override for ``/word/document.xml``
from the template content type to the document content type, and re-zips the
untouched parts (styles, numbering, header/footer, theme, media, ...) into
``colvir.docx``. Every style, list definition, and header/footer survives
unchanged; only the package-level content-type declaration is corrected.

Usage:
    python tools/build_default_docx_template.py

Repeatable: running this script again produces a functionally identical
``colvir.docx`` from the same ``colvir_v3.dotx`` source (same parts, same
bytes per part -- only the top-level zip member order/metadata could differ
across zippers, which is why callers should compare XML content, not raw
zip bytes, when verifying this asset).
"""

from __future__ import annotations

import zipfile
from pathlib import Path

_TOOLS_DIR = Path(__file__).resolve().parent
_TEMPLATES_DIR = _TOOLS_DIR.parent / "src" / "docgen" / "formatting" / "templates"
# The .dotx source is a build-time-only input -- unlike the derived
# colvir.docx it produces, it is never read at runtime (DocxExporter only
# ever opens colvir.docx) and does not belong in the runtime catalog
# directory or the packaged wheel, so it lives alongside this script
# instead of under formatting/templates.
_SOURCE_DOTX = _TOOLS_DIR / "colvir_v3.dotx"
_OUTPUT_DOCX = _TEMPLATES_DIR / "colvir.docx"

_CONTENT_TYPES_PART = "[Content_Types].xml"
_TEMPLATE_MAIN_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
)
_DOCUMENT_MAIN_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
)


def build(source: Path = _SOURCE_DOTX, output: Path = _OUTPUT_DOCX) -> Path:
    """Repackage `source` (.dotx) as a python-docx-openable `output` (.docx).

    Copies every zip part unchanged except `[Content_Types].xml`, where the
    Override for `/word/document.xml` is switched from the *template* main
    content type to the *document* main content type. Raises ValueError if
    the source does not declare the expected template content type (so a
    future template swap that already ships a document-typed part fails
    loudly instead of silently no-op'ing).
    """
    if not source.is_file():
        raise FileNotFoundError(f"Исходный шаблон не найден: {source}")

    with zipfile.ZipFile(source, "r") as archive_in:
        members = archive_in.infolist()
        content_types_xml = archive_in.read(_CONTENT_TYPES_PART).decode("utf-8")

        if _TEMPLATE_MAIN_CONTENT_TYPE not in content_types_xml:
            raise ValueError(
                f"Ожидаемый тип содержимого не найден в {_CONTENT_TYPES_PART}: "
                f"{_TEMPLATE_MAIN_CONTENT_TYPE}"
            )

        patched_content_types_xml = content_types_xml.replace(
            _TEMPLATE_MAIN_CONTENT_TYPE, _DOCUMENT_MAIN_CONTENT_TYPE
        )

        output.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive_out:
            for member in members:
                data = archive_in.read(member.filename)
                if member.filename == _CONTENT_TYPES_PART:
                    data = patched_content_types_xml.encode("utf-8")
                archive_out.writestr(member, data)

    return output


def main() -> None:
    output = build()
    # Verify the result actually opens with python-docx before declaring success.
    import docx

    document = docx.Document(str(output))
    print(f"Собран {output} ({len(document.paragraphs)} абзацев в теле шаблона).")


if __name__ == "__main__":
    main()
